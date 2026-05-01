# -*- coding: utf-8 -*-
"""
超无穹AI平台培训课程 - Word文档生成
基于python-docx库
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = r"D:\网站部署\超无穹项目\chaowuqiong-project\docs\培训课程\文档生成\word"

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_with_style(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading

def create_curriculum_document():
    """创建课程大纲Word文档"""
    doc = Document()

    doc.sections[0].page_width = Cm(29.7)
    doc.sections[0].page_height = Cm(21)
    doc.sections[0].left_margin = Cm(2)
    doc.sections[0].right_margin = Cm(2)
    doc.sections[0].top_margin = Cm(2)
    doc.sections[0].bottom_margin = Cm(2)

    title = doc.add_heading('超无穹AI平台培训课程大纲', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(22)

    doc.add_paragraph('培训周期：12周  |  总课时：72课时  |  目标：零基础到全栈工程师')
    doc.add_paragraph()

    phases = [
        {
            'name': '第一阶段：入门篇 (第1-2周)',
            'color': '4472C4',
            'description': '掌握计算机基础、HTML/CSS、Web开发入门',
            'hours': '12课时',
            'content': [
                ('计算机基础与开发环境', ['计算机基本操作', '认识互联网', '编程是什么', '开发工具介绍', '超无穹平台介绍']),
                ('Web基础入门', ['HTML初识', 'CSS入门', '响应式设计基础', '浏览器开发者工具']),
            ]
        },
        {
            'name': '第二阶段：基础篇 (第3-5周)',
            'color': '70AD47',
            'description': '掌握JavaScript、Node.js、数据库基础',
            'hours': '18课时',
            'content': [
                ('JavaScript基础', ['变量与数据类型', '运算符与表达式', '条件语句', '循环结构', '函数基础', '数组基础', '对象基础']),
                ('Node.js与后端基础', ['Node.js介绍', 'Express框架', '数据库概述', 'SQL基础', 'API设计入门']),
            ]
        },
        {
            'name': '第三阶段：进阶篇 (第6-8周)',
            'color': 'ED7D31',
            'description': '掌握React框架、平台功能开发',
            'hours': '18课时',
            'content': [
                ('React框架入门', ['React概述', '组件与Props', 'State与生命周期', '事件处理', '条件渲染与列表', 'Hooks深入']),
                ('超无穹平台开发', ['项目结构分析', 'AI对话功能', '角色管理', '会员系统', '移动端适配']),
                ('后端API开发', ['认证中间件', 'LLM配置模块', '支付集成', 'Redis缓存']),
            ]
        },
        {
            'name': '第四阶段：高阶篇 (第9-10周)',
            'color': '7030A0',
            'description': '掌握TypeScript、工程化、部署运维',
            'hours': '12课时',
            'content': [
                ('TypeScript与工程化', ['TypeScript入门', '状态管理', '前端路由', '构建工具', 'Git工作流']),
                ('部署与运维', ['服务器基础', 'Nginx配置', 'PM2进程管理', 'Docker基础', '监控与日志']),
            ]
        },
        {
            'name': '第五阶段：实战篇 (第11-12周)',
            'color': 'C00000',
            'description': '综合项目实战、答辩、就业指导',
            'hours': '12课时',
            'content': [
                ('综合项目实战', ['项目选题', '系统设计', '前端开发', '后端开发', '测试与修复', '部署上线']),
                ('结业与就业指导', ['项目答辩', '课程总结', '简历制作', '面试指导', '职业规划']),
            ]
        }
    ]

    for phase in phases:
        heading = doc.add_heading(phase['name'], level=1)
        for run in heading.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.color.rgb = RGBColor(int(phase['color'][:2], 16), int(phase['color'][2:4], 16), int(phase['color'][4:], 16))

        p = doc.add_paragraph()
        p.add_run(f"课时：{phase['hours']} | ").bold = True
        p.add_run(phase['description'])

        for module_name, topics in phase['content']:
            doc.add_heading(module_name, level=2)
            for topic in topics:
                doc.add_paragraph(topic, style='List Bullet')

        doc.add_paragraph()

    doc.save(os.path.join(OUTPUT_DIR, '01_课程大纲.docx'))
    print(f"✅ 已生成：01_课程大纲.docx")

def create_teaching_plan_document():
    """创建教学计划Word文档"""
    doc = Document()

    doc.add_heading('超无穹AI平台培训 - 教学计划', 0)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('一、课程基本信息')
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('课程名称', '超无穹AI平台全栈开发工程师培训'),
        ('培训周期', '12周（约3个月）'),
        ('总课时', '72课时（每天1课时，每周6天）'),
        ('班级规模', '25-30人/班'),
        ('教学方法', '理论+实操+项目驱动'),
        ('考核方式', '出勤+作业+阶段测试+实践项目+结业答辩'),
    ]
    for i, (key, val) in enumerate(info_data):
        info_table.rows[i].cells[0].text = key
        info_table.rows[i].cells[1].text = val
        set_cell_shading(info_table.rows[i].cells[0], 'D9E2F3')

    doc.add_paragraph()
    doc.add_paragraph('二、周教学进度计划')

    weeks = [
        ('第1周', '计算机基础与开发环境', '计算机操作、互联网基础、编程概念、开发工具'),
        ('第2周', 'Web基础入门', 'HTML标签、CSS样式、响应式设计、开发者工具'),
        ('第3周', 'JavaScript基础（上）', '变量、数据类型、运算符、条件语句'),
        ('第4周', 'JavaScript基础（下）', '循环、函数、数组、对象、DOM操作'),
        ('第5周', 'Node.js与数据库', 'Node.js、Express、MySQL、SQL基础'),
        ('第6周', 'React框架入门', '组件、Props、State、Hooks'),
        ('第7周', '平台前端开发', 'AI对话、角色管理、会员系统'),
        ('第8周', '后端API开发', '认证、JWT、支付集成、Redis'),
        ('第9周', 'TypeScript与工程化', 'TS类型、状态管理、Git工作流'),
        ('第10周', '部署与运维', 'Nginx、Docker、PM2、监控'),
        ('第11周', '综合项目实战', '项目开发、测试、部署'),
        ('第12周', '答辩与就业', '项目答辩、简历指导、面试培训'),
    ]

    plan_table = doc.add_table(rows=len(weeks)+1, cols=3)
    plan_table.style = 'Table Grid'
    headers = ['周次', '主题', '主要内容']
    for i, h in enumerate(headers):
        plan_table.rows[0].cells[i].text = h
        set_cell_shading(plan_table.rows[0].cells[i], '4472C4')
        plan_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (week, theme, content) in enumerate(weeks):
        plan_table.rows[i+1].cells[0].text = week
        plan_table.rows[i+1].cells[1].text = theme
        plan_table.rows[i+1].cells[2].text = content

    doc.save(os.path.join(OUTPUT_DIR, '02_教学计划.docx'))
    print(f"✅ 已生成：02_教学计划.docx")

def create_practice_document():
    """创建实践项目Word文档"""
    doc = Document()

    doc.add_heading('超无穹AI平台培训 - 实践项目手册', 0)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('本课程包含9个实践项目，循序渐进，从简单到复杂。')

    projects = [
        ('P1', '个人简历网站', 'HTML/CSS', '⭐', '第一阶段', '运用HTML/CSS制作个人简历页面，掌握页面结构与样式'),
        ('P2', '待办事项应用', 'JavaScript', '⭐⭐', '第二阶段', '使用原生JS实现Todo App，掌握DOM操作与事件处理'),
        ('P3', '用户管理系统', 'Node.js/MySQL', '⭐⭐', '第二阶段', '实现用户注册、登录、资料管理功能'),
        ('P4', 'AI对话组件', 'React', '⭐⭐⭐', '第三阶段', '开发超无穹平台的AI对话界面组件'),
        ('P5', '会员中心', 'React/Node.js', '⭐⭐⭐', '第三阶段', '实现会员套餐展示、购买、积分管理'),
        ('P6', '支付功能', '支付API', '⭐⭐⭐', '第三阶段', '对接支付宝/微信支付功能'),
        ('P7', 'TypeScript重构', 'TypeScript', '⭐⭐⭐', '第四阶段', '将项目改造为TypeScript，提升类型安全'),
        ('P8', '容器化部署', 'Docker/Nginx', '⭐⭐⭐⭐', '第四阶段', '使用Docker容器化部署整个平台'),
        ('P9', '结业综合项目', '全栈技术', '⭐⭐⭐⭐⭐', '第五阶段', '完成完整的全栈项目并部署上线'),
    ]

    table = doc.add_table(rows=len(projects)+1, cols=6)
    table.style = 'Table Grid'
    headers = ['编号', '项目名称', '技术栈', '难度', '阶段', '描述']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], '4472C4')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, p in enumerate(projects):
        for j, val in enumerate(p):
            table.rows[i+1].cells[j].text = val

    doc.save(os.path.join(OUTPUT_DIR, '03_实践项目手册.docx'))
    print(f"✅ 已生成：03_实践项目手册.docx")

def create_assessment_document():
    """创建考核标准Word文档"""
    doc = Document()

    doc.add_heading('超无穹AI平台培训 - 考核标准', 0)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('一、毕业要求')
    req_table = doc.add_table(rows=6, cols=3)
    req_table.style = 'Table Grid'
    req_data = [
        ('要求项', '标准', '权重'),
        ('出勤率', '≥ 90%', '10%'),
        ('作业完成', '全部提交 + 及格以上', '20%'),
        ('阶段测试', '每阶段≥60分', '20%'),
        ('实践项目', '完成全部9个实践项目', '30%'),
        ('结业答辩', '评委会≥60分通过', '20%'),
    ]
    for i, row in enumerate(req_data):
        for j, val in enumerate(row):
            req_table.rows[i].cells[j].text = val
            if i == 0:
                set_cell_shading(req_table.rows[i].cells[j], '4472C4')
                req_table.rows[i].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    doc.add_paragraph('二、能力等级标准')

    levels = [
        ('L1-入门', '第一阶段', 'HTML/CSS/JS基础', '能写静态页面'),
        ('L2-初级', '第二阶段', '前端开发/后端基础', '能做简单功能'),
        ('L3-中级', '第三阶段', 'React/Node.js/数据库', '能独立开发模块'),
        ('L4-高级', '第四阶段', '工程化/部署/性能优化', '能做性能调优'),
        ('L5-结业', '第五阶段', '全栈能力/团队协作', '能完成完整项目'),
    ]

    level_table = doc.add_table(rows=len(levels)+1, cols=4)
    level_table.style = 'Table Grid'
    level_headers = ['能力等级', '对应阶段', '技术能力', '项目能力']
    for i, h in enumerate(level_headers):
        level_table.rows[0].cells[i].text = h
        set_cell_shading(level_table.rows[0].cells[i], '70AD47')
        level_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, row in enumerate(levels):
        for j, val in enumerate(row):
            level_table.rows[i+1].cells[j].text = val

    doc.save(os.path.join(OUTPUT_DIR, '04_考核标准.docx'))
    print(f"✅ 已生成：04_考核标准.docx")

def main():
    print("=" * 50)
    print("超无穹AI平台培训课程 - Word文档生成")
    print("=" * 50)

    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

    create_curriculum_document()
    create_teaching_plan_document()
    create_practice_document()
    create_assessment_document()

    print("=" * 50)
    print(f"所有Word文档已生成到：{OUTPUT_DIR}")
    print("=" * 50)

if __name__ == '__main__':
    main()
